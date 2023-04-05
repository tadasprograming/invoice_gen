from docx import Document
from scr import get_data
#import scrape_rekvizitai

#Pirma surenkam visus duomenis reikalingus sąskaitai

invoice_date = input('Įveskite sąskaitos datą (yyyy-mm-dd): ')
invoice_number = invoice_date + '-01'

#Klausiam ar nori scrapinti ar vesti ranka, abiem atvejais gaunam tuple (pavadinimas,kodas,adresas,vadovas)

if_scrape = ''
while if_scrape != 't' and if_scrape != 'n':
    if_scrape = input('Ar norite įmonės duomenis bandyti nuskaityti įš rekvizitai.vz.lt? (t ar n): ')

if if_scrape == 't':
    company_data = get_data()
else:
    print('Įveskite įmonės duomenis')
    company_name = input('Įveskite įmonės pavadinimą: ')
    company_code = input('Įveskite įmonės kodą: ')
    company_address = input('Įveskite įmonės adresą: ')
    company_boss = input('Įveskite įmonės vadovo informaciją: ')
    company_data = (company_name,company_code,company_address,company_boss)

#Prašom likusių duomenų, kuriuos reikia suvesti ranka

services = input('Įveskite suteiktų paslaugų pavadinimą: ')
amount = input('Įveskite paslaugų kiekį: ')
unit_price = input('Įveskite paslaugų vieneto kainą: ')
sum = int(amount) * float(unit_price)
print(f'Bendra suma {sum}')
sum_lt = input('Įveskite sumą žodžiais LT: ')
sum_en = input('Įveskite sumą žodžiais ENG: ')

#Atsidarom template failą ir surašom visus duomenis į failą    

template = Document('invoice_template.docx')

#Įrašom sąskaitos numerį ir datą

for paragraph in template.paragraphs:
    if '{invoice number}' in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if '{invoice number}' in inline[i].text:
                inline[i].text = inline[i].text.replace('{invoice number}', invoice_number)

for paragraph in template.paragraphs:
    if '{invoice date}' in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if '{invoice date}' in inline[i].text:
                inline[i].text = inline[i].text.replace('{invoice date}', invoice_date)

#Įrašom įmonės informaciją į 0 table, 0 row, 1 collum 

company_info = {'company_name':{'from':'{name}','to':company_data[0]},
                'company_code':{'from':'{code}','to':company_data[1]},
                'company_address':{'from':'{address}','to':company_data[2]},
                'company_boss':{'from':'{boss}','to':company_data[3]}
                }

table = template.tables[0]
cell = table.cell(0,1)

for key in company_info:
    for paragraph in cell.paragraphs:
        if company_info[key]['from'] in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if company_info[key]['from'] in inline[i].text:
                    inline[i].text = inline[i].text.replace(company_info[key]['from'], company_info[key]['to'])

#Susikuriam funciją įrašyti duomenis į tam tikrą lentelės langelį:

def write_cell(info,key_word,table_no,row_no,cell_no):

    table = template.tables[table_no]
    cell = table.cell(row_no,cell_no)
    for paragraph in cell.paragraphs:
        if key_word in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if key_word in inline[i].text:
                    inline[i].text = inline[i].text.replace(key_word, info)

#Įrašom likusius duomenis į lentelę

write_cell(services,'{services}',1,1,1)
write_cell(amount,'{amount}',1,1,3)
write_cell(unit_price,'{unit price}',1,1,4)
sum = str(sum)
write_cell(sum,'{sum}',1,1,5)
write_cell(sum,'{total sum}',1,2,1)
write_cell(sum_lt,'{sum in words LT}',2,0,0)
write_cell(sum_en,'{sum in words ENG}',2,0,0)

#Sukuriam sąskaitai pavadinimą ir išsaugom sąskaitą

formated_date = invoice_date.replace(' ','')
invoice_name = 'TVF '+invoice_date+'_01_INVOICE.docx'
template.save(invoice_name)
print('Sąskaita sėkmingai sukurta!')