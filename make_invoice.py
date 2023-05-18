from docx import Document

TEMPLATES = {"Tado": 'tadas_invoice_template.docx',
             "Pamelos": 'pamela_invoice_template.docx'}

INITIALS = {"Tado": 'TVF', "Pamelos": 'PG'}


def make_invoice(invoice_info, company_info, user):
    invoice = Document(TEMPLATES[user])

    # Įrašom sąskaitos numerį ir datą

    for paragraph in invoice.paragraphs:
        if '{invoice number}' in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '{invoice number}' in inline[i].text:
                    inline[i].text = inline[i].text.replace(
                        '{invoice number}', invoice_info['invoice_number'])

    for paragraph in invoice.paragraphs:
        if '{invoice date}' in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '{invoice date}' in inline[i].text:
                    inline[i].text = inline[i].text.replace(
                        '{invoice date}', invoice_info['invoice_date'])

    # Įrašom įmonės informaciją į 0 table, 0 row, 1 collum

    company = {
        'company_name': {'from': '{name}', 'to': company_info['name']},
        'company_code': {'from': '{code}', 'to': company_info['code']},
        'company_address': {'from': '{address}',
                            'to': company_info['address']},
        'company_boss': {'from': '{boss}', 'to': company_info['boss']},
        'company_extra_info': {'from': '{extra_info}',
                               'to': company_info['extra_info']}}

    table = invoice.tables[0]
    cell = table.cell(0, 1)

    for key in company:
        for paragraph in cell.paragraphs:
            if company[key]['from'] in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if company[key]['from'] in inline[i].text:
                        inline[i].text = inline[i].text.replace(
                            company[key]['from'], company[key]['to'])

    # Susikuriam funciją įrašyti duomenis į tam tikrą lentelės langelį:

    def write_cell(info, key_word, table_no, row_no, cell_no):

        table = invoice.tables[table_no]
        cell = table.cell(row_no, cell_no)
        for paragraph in cell.paragraphs:
            if key_word in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if key_word in inline[i].text:
                        inline[i].text = inline[i].text.replace(key_word, info)

    # Įrašom likusius duomenis į lentelę

    write_cell(invoice_info['services'], '{services}', 1, 1, 1)
    write_cell(invoice_info['amount'], '{amount}', 1, 1, 3)
    write_cell(invoice_info['unit_price'], '{unit price}', 1, 1, 4)
    write_cell(invoice_info['sum'], '{sum}', 1, 1, 5)
    write_cell(invoice_info['sum'], '{total sum}', 1, 2, 5)
    write_cell(invoice_info['total_LT'], '{sum in words LT}', 2, 0, 0)
    write_cell(invoice_info['total_EN'], '{sum in words ENG}', 2, 0, 0)

    # Sukuriam sąskaitai pavadinimą ir išsaugom sąskaitą

    formated_date = invoice_info['invoice_date'].replace('-', '')
    invoice_name = f'{INITIALS[user]}_{formated_date}_01_INVOICE.docx'
    invoice.save(invoice_name)
